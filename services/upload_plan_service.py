import os
import re
import json
import asyncio
from typing import Dict, Any, Optional, List, Tuple
import PyPDF2
import docx
from docx import Document
import tempfile
from openai import AsyncOpenAI

client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# Model used for every extraction call in this pipeline. Configurable via env
# so the model can be rolled forward without touching code.
#
# Default is `gpt-4o`, which is the most-capable model guaranteed to be
# enabled on every OpenAI account today AND supports the strict JSON Schema
# structured-output mode this extractor relies on. To roll forward to gpt-5
# (or any successor) set `OPENAI_EXTRACTION_MODEL=gpt-5` in the backend env;
# no code change required. If the configured model is not enabled on the
# account, `extract_per_question_answers` will fail loudly so the
# misconfiguration surfaces immediately instead of producing silent 0%
# extraction results.
_EXTRACTION_MODEL = os.getenv("OPENAI_EXTRACTION_MODEL", "gpt-4o")

# Input cap. gpt-5 / gpt-4o have a 128k+ token window; 150k chars (~37k tokens)
# covers a 70-page business plan with room left for the questions list, schema,
# and JSON output. Larger than this almost certainly indicates a non-plan upload.
_EXTRACTION_MAX_CHARS = 150_000


_CANONICAL_BP_QUESTIONS_CACHE: Optional[List[Tuple[int, str]]] = None


def parse_canonical_business_plan_questions() -> List[Tuple[int, str]]:
    """
    Parse the canonical 45 BUSINESS_PLAN questions from ANGEL_SYSTEM_PROMPT in
    constant.py — the single source of truth used by Angel itself. Returns a sorted
    list of (question_number, first_line_question_text) tuples. Cached after first
    parse since the prompt is static.

    The regex is anchored to start-of-line (^ with re.MULTILINE) because canonical
    definitions appear at column 0 in the prompt:
        [[Q:BUSINESS_PLAN.08]] Who is your target customer?...
    while inline mentions (which would otherwise hijack the capture) are always
    embedded mid-sentence:
        Market Research applies to [[Q:BUSINESS_PLAN.08]] through [[Q:BUSINESS_PLAN.13]].
    Without the anchor, the regex grabbed "through" as Q8's text, etc.
    """
    global _CANONICAL_BP_QUESTIONS_CACHE
    if _CANONICAL_BP_QUESTIONS_CACHE is not None:
        return _CANONICAL_BP_QUESTIONS_CACHE

    from utils.constant import ANGEL_SYSTEM_PROMPT

    # ^ + MULTILINE → only definitions at the start of a line are captured.
    pattern = re.compile(
        r'^\[\[Q:BUSINESS_PLAN\.(\d{2})\]\]\s*([^\n\[]+)',
        re.MULTILINE,
    )
    seen: Dict[int, str] = {}
    for num_str, text in pattern.findall(ANGEL_SYSTEM_PROMPT):
        num = int(num_str)
        if not (1 <= num <= 45):
            continue
        cleaned = text.strip()
        if cleaned and num not in seen:
            seen[num] = cleaned

    result = sorted(seen.items())
    _CANONICAL_BP_QUESTIONS_CACHE = result
    return result


def _group_canonical_questions(
    questions: List[Tuple[int, str]],
) -> List[Tuple[str, List[Tuple[int, str]]]]:
    """Group the 45 canonical questions into the 8 UI categories used by the
    analysis modal. Group ordering matches `_category_for_question` so the
    extraction pipeline maps 1:1 to what the user sees in the modal."""
    buckets: Dict[str, List[Tuple[int, str]]] = {
        "Business Overview": [],
        "Market & Customers": [],
        "Operations": [],
        "Brand & Marketing": [],
        "Legal & Regulatory": [],
        "Financials": [],
        "Growth & Long-Term": [],
        "Risk & Vision": [],
    }
    for num, text in questions:
        buckets[_category_for_question(num)].append((num, text))
    return [(name, qs) for name, qs in buckets.items() if qs]


def _build_extraction_schema(group_questions: List[Tuple[int, str]]) -> Dict[str, Any]:
    """Build an OpenAI structured-output JSON Schema for a single extraction
    group. Strict mode forces the API to return every question number as a key
    with either a string answer or null — no missing keys, no extra keys, no
    type ambiguity to parse around on the Python side."""
    keys = [str(num) for num, _ in group_questions]
    answer_properties = {
        str(num): {
            "anyOf": [{"type": "string"}, {"type": "null"}],
            "description": text,
        }
        for num, text in group_questions
    }
    return {
        "name": "extracted_answers",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "answers": {
                    "type": "object",
                    "properties": answer_properties,
                    "required": keys,
                    "additionalProperties": False,
                },
            },
            "required": ["answers"],
            "additionalProperties": False,
        },
    }


_EXTRACTION_SYSTEM_PROMPT = (
    "You extract structured answers from business plan documents. "
    "Your job is recall: if the document contains information that answers a "
    "question — anywhere, in any wording, in narrative paragraphs OR in table "
    "cells — you must extract it. Only return null when the document genuinely "
    "has no information on that topic. Never invent. Return strict JSON only."
)


async def _extract_group(
    group_name: str,
    group_questions: List[Tuple[int, str]],
    document_text: str,
    model: str,
) -> Dict[int, Optional[str]]:
    """Run extraction for one category. Returns {q_num: answer or None} for
    exactly the questions in this group. The full document is provided to every
    group call so cross-section evidence (e.g. legal structure mentioned inside
    the Operations table) is never missed."""
    questions_block = "\n".join(f"Q{num}. {text}" for num, text in group_questions)

    user_prompt = f"""CATEGORY: {group_name}

QUESTIONS TO ANSWER:
{questions_block}

DOCUMENT (full text — narrative sections AND table rows):
{document_text}

RULES:
1. Read the ENTIRE document. Answers often live inside tables — Company Overview, Financial Projections, Risk Matrix, Implementation Timeline, Marketing Channels, Operations Structure, etc. — and not only in narrative paragraphs. Treat table cells as a primary source.
2. The document may phrase things differently than the question. Match by meaning, not by literal wording (e.g. "Revenue Streams: Subscription fees" answers "How will your business make money?"; "Funding Requirements: $350,000, Personal savings, angel investors, venture capital" answers both the funding-source and total-cost questions).
3. EXPLICIT "NOT YET" ENTRIES ARE VALID ANSWERS, NOT NULLS. If the document says things like "Not yet specified", "Not yet decided", "TBD", "To be determined", "Currently in the idea stage", "Not yet defined", etc. for a topic — that is the founder's current position on that question. Return it verbatim (or paraphrased: "Not yet decided — the founder has not chosen an X yet."). Do NOT return null for these.
4. For each question, write a 1-4 sentence answer in the founder's voice using only information present in the document.
5. Only return null when the document genuinely contains nothing on that question's topic — not present, not even as "TBD". Wording mismatch is never a reason to return null.
6. Do not include the question text in the answer. Do not invent details that are not in the document.

Respond with strict JSON matching the provided schema."""

    response = await client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": _EXTRACTION_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
        seed=1,
        response_format={"type": "json_schema", "json_schema": _build_extraction_schema(group_questions)},
    )

    parsed = json.loads(response.choices[0].message.content)
    answers_raw = parsed["answers"]

    out: Dict[int, Optional[str]] = {}
    for num, _ in group_questions:
        value = answers_raw.get(str(num))
        if isinstance(value, str):
            cleaned = value.strip()
            out[num] = cleaned if cleaned and cleaned.upper() not in {"NULL", "N/A", "NOT_FOUND"} else None
        else:
            # Strict-schema guarantees string|null, so anything else means the
            # value was the literal JSON null — surface as Python None.
            out[num] = None
    return out


async def extract_per_question_answers(content: str) -> Dict[int, Optional[str]]:
    """Source-of-truth extraction for the upload pipeline.

    Architecture
    ------------
    The canonical 45 questions are split into 8 category groups (the same
    groups the analysis modal renders), and each group runs as a separate
    LLM call with the full document attached. All 8 calls execute in
    parallel via asyncio.gather, so wall-clock time matches a single call
    while each call only has to focus on 5-7 closely-related questions.

    Why per-category instead of one giant call
    ------------------------------------------
    A single 45-question call was missing answers buried in later sections
    of long documents — the model triaged attention across too many disparate
    topics and output-token pressure clipped late answers. With 8 focused
    calls, each call has plenty of output budget and a coherent topic to
    nail. Cost is essentially flat (input dominates and is the same total
    text); recall is dramatically higher.

    Determinism
    -----------
    Every call uses temperature=0 and a fixed seed, plus OpenAI's strict
    JSON Schema structured outputs. Same document in → same answers out.
    """
    questions = parse_canonical_business_plan_questions()
    if not questions:
        return {}

    truncated_content = content[:_EXTRACTION_MAX_CHARS]
    truncation_notice = (
        f"\n\n[Document is {len(content):,} chars; first {_EXTRACTION_MAX_CHARS:,} included.]"
        if len(content) > _EXTRACTION_MAX_CHARS
        else ""
    )
    document_text = truncated_content + truncation_notice

    groups = _group_canonical_questions(questions)
    print(
        f"📑 Extracting {len(questions)} questions in {len(groups)} parallel groups "
        f"using {_EXTRACTION_MODEL} ({len(document_text):,} chars of document)"
    )

    tasks = [
        _extract_group(group_name, group_qs, document_text, _EXTRACTION_MODEL)
        for group_name, group_qs in groups
    ]
    group_results = await asyncio.gather(*tasks, return_exceptions=True)

    merged: Dict[int, Optional[str]] = {num: None for num, _ in questions}
    failures: List[Tuple[str, BaseException]] = []
    successes = 0
    for (group_name, group_qs), result in zip(groups, group_results):
        if isinstance(result, BaseException):
            failures.append((group_name, result))
            print(f"⚠️ Extraction group '{group_name}' failed: {type(result).__name__}: {result}")
            continue
        successes += 1
        for num in (n for n, _ in group_qs):
            merged[num] = result.get(num)

    # A single transient failure must not erase the other groups' answers —
    # those questions just stay null and the analysis modal renders them as
    # missing. But if EVERY group failed, the cause is systemic (bad model
    # name, no API key, rate limit, account out of credits, etc.) and the
    # right behaviour is to raise so the upload endpoint returns a real 500
    # with the actual OpenAI error. Returning all-null silently turns this
    # into "extractor doesn't work" with no diagnosis path for the user.
    if successes == 0 and failures:
        first_group, first_exc = failures[0]
        raise RuntimeError(
            f"All {len(groups)} extraction groups failed using model "
            f"{_EXTRACTION_MODEL!r}. First error from '{first_group}': "
            f"{type(first_exc).__name__}: {first_exc}"
        ) from first_exc

    filled = sum(1 for v in merged.values() if v)
    print(
        f"📑 Extraction complete: {filled}/{len(questions)} answered "
        f"({successes}/{len(groups)} groups succeeded)"
    )
    return merged


# Maps canonical question numbers (1–45) to legacy business_info summary keys that
# downstream code (business_context, prompts, frontend) already reads. This lets us
# populate the summary dict without breaking existing consumers.
_QUESTION_TO_SUMMARY_KEYS: Dict[int, List[str]] = {
    1: ["business_idea", "solution"],
    2: ["product_or_service", "product_description"],
    3: ["unique_advantage", "competitive_advantage", "value_proposition", "unique_value"],
    4: ["business_stage"],
    5: ["business_name"],
    6: ["industry", "business_type"],
    7: ["short_term_goals"],
    8: ["target_market"],
    9: ["distribution_channels"],
    10: ["problem_solved", "problem"],
    11: ["competitors"],
    13: ["differentiation"],
    14: ["location"],
    18: ["mission", "tagline", "vision"],
    19: ["marketing_methods", "marketing_strategy"],
    21: ["unique_selling_proposition"],
    24: ["legal_structure", "business_structure"],
    29: ["revenue_model"],
    30: ["pricing", "pricing_strategy"],
    32: ["funding_source"],
    33: ["first_year_financial_goals", "key_metrics"],
    34: ["startup_costs"],
    36: ["long_term_goals", "goals"],
    44: ["funding_needs"],
    45: ["five_year_vision"],
}


def _summary_from_per_question(per_question: Dict[int, Optional[str]]) -> Dict[str, Any]:
    """Build the legacy business_info summary dict (preserving downstream key names)
    from the canonical per-question answers."""
    summary: Dict[str, Any] = {}
    for q_num, keys in _QUESTION_TO_SUMMARY_KEYS.items():
        value = per_question.get(q_num)
        for key in keys:
            summary[key] = value
    return summary

async def process_uploaded_plan(file_path: str, file_extension: str) -> str:
    """
    Process uploaded business plan file and extract text content
    Supports: PDF, DOC, DOCX, TXT
    """
    try:
        if file_extension == '.pdf':
            return await extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return await extract_docx_text(file_path)
        elif file_extension == '.doc':
            return await extract_doc_text(file_path)
        elif file_extension == '.txt':
            return await extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        print(f"Error processing file: {e}")
        raise Exception(f"Failed to process file: {str(e)}")

async def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")

async def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")

async def extract_doc_text(file_path: str) -> str:
    """Extract text from DOC file (requires additional library)"""
    try:
        # For .doc files, we'd need python-docx2txt or similar
        # For now, return a message to convert to .docx
        raise Exception("DOC files are not supported. Please convert to DOCX format.")
    except Exception as e:
        raise Exception(f"Error extracting DOC text: {str(e)}")

async def extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        raise Exception(f"Error extracting TXT text: {str(e)}")

async def extract_business_info_from_plan(
    content: str,
    per_question_answers: Optional[Dict[int, Optional[str]]] = None,
) -> Dict[str, Any]:
    """
    Build the legacy business_info summary dict from canonical per-question answers.

    Downstream code (business_context fields, generate_plan_service, prompts in
    angel_service, frontend display) reads keys like business_name, industry,
    mission, location, target_market — those are preserved by mapping canonical
    question answers back to those keys via _QUESTION_TO_SUMMARY_KEYS.

    `per_question_answers` is normally provided by the caller (already extracted
    from the document) so we don't double-call the LLM. When omitted, this helper
    extracts on demand for backward compatibility.

    When extraction yields nothing, the summary is empty — that is the correct
    signal. We do NOT fabricate a fake summary from keyword regexes; an "all
    empty" result from the 8-group structured extractor means the document
    genuinely isn't a business plan, and the upload modal should say so.
    """
    if per_question_answers is None:
        per_question_answers = await extract_per_question_answers(content)

    return _summary_from_per_question(per_question_answers)

async def validate_business_plan_content(content: str) -> Dict[str, Any]:
    """
    Validate that the uploaded content is actually a business plan
    """
    try:
        validation_prompt = f"""
        Analyze this document and determine if it's a business plan. Return JSON with:
        {{
            "is_business_plan": true/false,
            "confidence": 0.0-1.0,
            "missing_sections": ["list of missing typical business plan sections"],
            "content_type": "description of what type of document this is",
            "recommendations": "suggestions for improvement"
        }}

        Document content:
        {content[:4000]}

        Typical business plan sections include: Executive Summary, Company Description, Market Analysis, Organization, Service/Product Line, Marketing, Financial Projections, etc.
        """

        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert at analyzing business documents and plans."},
                {"role": "user", "content": validation_prompt}
            ],
            temperature=0.3,
            max_tokens=1000
        )

        validation_result = response.choices[0].message.content.strip()
        
        if validation_result.startswith('```json'):
            validation_result = validation_result[7:]
        if validation_result.endswith('```'):
            validation_result = validation_result[:-3]
            
        return json.loads(validation_result)
        
    except Exception as e:
        print(f"Error validating business plan: {e}")
        return {
            "is_business_plan": True,  # Default to accepting
            "confidence": 0.5,
            "missing_sections": [],
            "content_type": "Unknown document type",
            "recommendations": "Unable to validate document structure"
        }

async def analyze_plan_completeness(
    content: str,
    business_info: Dict[str, Any],
    per_question_answers: Optional[Dict[int, Optional[str]]] = None,
) -> Dict[str, Any]:
    """
    Build the analysis payload (summary, completeness score, missing_questions list)
    directly from canonical per-question answers. The previous version compared the
    document against a hardcoded 45-question list that had drifted out of sync with
    Angel's actual questionnaire — the canonical source is now ANGEL_SYSTEM_PROMPT.
    """
    canonical = parse_canonical_business_plan_questions()

    if per_question_answers is None:
        per_question_answers = await extract_per_question_answers(content)

    if not canonical:
        return {
            "summary": "Unable to load canonical business plan questions.",
            "completeness_score": 0.0,
            "found_information": {},
            "missing_questions": [],
            "recommendations": "An internal error occurred. Please proceed manually.",
        }

    found_question_numbers: List[int] = []
    missing_question_entries: List[Dict[str, Any]] = []
    for q_num, q_text in canonical:
        answer = per_question_answers.get(q_num)
        if answer and answer.strip():
            found_question_numbers.append(q_num)
        else:
            missing_question_entries.append({
                "question_number": q_num,
                "question_text": q_text,
                "category": _category_for_question(q_num),
                "priority": _priority_for_question(q_num),
            })

    completeness_score = round(len(found_question_numbers) / len(canonical), 2)

    # Each summary flag is anchored to the SINGLE primary canonical question
    # for that topic. The previous version OR'd in adjacent questions as a
    # forgiveness layer for inconsistent single-call extraction; with the new
    # per-category strict-schema extractor that hedge is no longer needed and
    # was actively misleading users — e.g. lighting up "Business Name ✓" in
    # the summary while showing "Q5 Business Name" in the missing list. The
    # summary now agrees with the missing list by construction.
    def _filled(num: int) -> bool:
        return bool(per_question_answers.get(num))

    found_information = {
        "business_name": _filled(5),
        "mission_vision": _filled(18),
        "problem_solution": _filled(10),
        "target_market": _filled(8),
        "competitors": _filled(11),
        "financial_projections": _filled(29),
        "marketing_strategy": _filled(19),
        "operational_plan": _filled(14),
        "legal_structure": _filled(24),
        "risk_analysis": _filled(42),
    }

    if completeness_score >= 0.85:
        summary_line = f"Your plan covers {len(found_question_numbers)} of {len(canonical)} Founderport questions. Only a few gaps remain."
    elif completeness_score >= 0.5:
        summary_line = f"Your plan covers {len(found_question_numbers)} of {len(canonical)} Founderport questions. We will fill the {len(missing_question_entries)} remaining gaps together."
    else:
        summary_line = f"Your plan covers {len(found_question_numbers)} of {len(canonical)} Founderport questions. Most sections still need detail — we will work through them step by step."

    return {
        "summary": summary_line,
        "completeness_score": completeness_score,
        "found_information": found_information,
        "found_question_numbers": found_question_numbers,
        "missing_questions": missing_question_entries,
        "recommendations": "Complete the remaining questions to produce your full Founderport business plan.",
    }


def _category_for_question(q_num: int) -> str:
    """Approximate section grouping for the analysis modal display."""
    if 1 <= q_num <= 7:
        return "Business Overview"
    if 8 <= q_num <= 13:
        return "Market & Customers"
    if 14 <= q_num <= 17:
        return "Operations"
    if 18 <= q_num <= 23:
        return "Brand & Marketing"
    if 24 <= q_num <= 28:
        return "Legal & Regulatory"
    if 29 <= q_num <= 34:
        return "Financials"
    if 35 <= q_num <= 41:
        return "Growth & Long-Term"
    if 42 <= q_num <= 45:
        return "Risk & Vision"
    return "General"


def _priority_for_question(q_num: int) -> str:
    """Coarse priority hint for the analysis modal — high for foundational items."""
    if q_num in {1, 2, 3, 5, 8, 10, 18, 24, 29}:
        return "high"
    if q_num in {6, 7, 11, 13, 14, 19, 21, 30, 32, 33, 42}:
        return "medium"
    return "low"


